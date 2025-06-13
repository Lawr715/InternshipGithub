import os
import re
import json
import csv
from io import StringIO, BytesIO
from datetime import datetime
from collections import Counter
import logging
import warnings
import seaborn as sns
from pathlib import Path
from docx import Document
import matplotlib.pyplot as plt
warnings.filterwarnings("ignore", message="Examining the path of torch.classes raised")

import streamlit as st

st.set_page_config(page_title="Transcript Hub", layout="wide")

# --- RESET LOGIC (must be before session state usage) ---
RESET_ON_START = True  # ← flip to False if you don't want to wipe every time

if RESET_ON_START and not st.session_state.get("has_been_reset", False):
    st.session_state.clear()
    # Optional: Delete all transcript files
    for f in Path("transcripts").glob("*.json"):
        f.unlink()
    st.session_state["has_been_reset"] = True

# --- File-based transcript storage ---
TRANSCRIPT_DIR = Path("transcripts")
TRANSCRIPT_DIR.mkdir(exist_ok=True)

def save_transcript_to_file(transcript):
    filename = f"{transcript['date']}_{transcript['type']}_{transcript['location'].replace(' ', '_')}.json"
    filepath = TRANSCRIPT_DIR / filename
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(transcript, f, indent=2, ensure_ascii=False)

def load_transcripts_from_files():
    transcripts = []
    for file in TRANSCRIPT_DIR.glob("*.json"):
        with open(file, "r", encoding="utf-8") as f:
            transcript = json.load(f)
            transcripts.append(transcript)
    return transcripts

def update_sidebar_filters():
    st.session_state["filter_sites"] = sorted({t["location"] for t in st.session_state["transcripts"] if t.get("location")})
    st.session_state["filter_types"] = sorted({t["type"] for t in st.session_state["transcripts"] if t.get("type")})
    st.session_state["filter_languages"] = sorted({t["language"] for t in st.session_state["transcripts"] if t.get("language")})

def delete_transcript_file(transcript):
    filename = f"{transcript['date']}_{transcript['type']}_{transcript['location'].replace(' ', '_')}.json"
    filepath = TRANSCRIPT_DIR / filename
    if filepath.exists():
        filepath.unlink()

# Helper to read .docx content
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_docx(transcript):
    doc = Document()
    doc.add_heading(f"{transcript['type']} – {transcript['location']} – {transcript['date']}", level=1)
    doc.add_paragraph("Summary:")
    doc.add_paragraph(transcript.get("summary", ""))
    doc.add_paragraph("Top Themes:")
    doc.add_paragraph(", ".join(transcript.get("themes", [])))
    doc.add_paragraph("Transcript:")
    doc.add_paragraph(transcript["text"])

    if transcript["tags"]:
        doc.add_paragraph("Tags: " + ", ".join(transcript["tags"]))
    if transcript["comments"]:
        doc.add_paragraph("Comments:")
        for comment in transcript["comments"]:
            doc.add_paragraph(f"- {comment}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if "transcripts" not in st.session_state:
    st.session_state["transcripts"] = load_transcripts_from_files()
    update_sidebar_filters()

# Ensure all transcripts have the required fields
for t in st.session_state["transcripts"]:
    t.setdefault("language", "Unknown")
    t.setdefault("tags", [])
    t.setdefault("comments", [])
    t.setdefault("themes", [])
    t.setdefault("summary", "")

PASSWORD = "transcripthub2025"

if not st.session_state.get("authenticated", False):
    st.title("🔐 Secure Access")
    st.markdown("Please enter the password to access **Transcript Hub**.")
    password_input = st.text_input("Password", type="password")

    if password_input == PASSWORD:
        st.session_state["authenticated"] = True
        st.success("Access granted! Loading app...")
        st.rerun()
    elif password_input:
        st.error("Incorrect password.")
    
    st.stop()  # Prevents the rest of the app from rendering

st.title("🧰 Transcript Hub")
st.caption("Internal tool to manage and analyze finalized KII and FGD transcripts")

# Filtering (move this above tab3 for use in all tabs)
# Always show filters — update dynamically when transcripts exist
sites = st.session_state.get("filter_sites", [])
types = st.session_state.get("filter_types", [])
languages = st.session_state.get("filter_languages", [])

selected_site = st.sidebar.selectbox("Site", ["All"] + sites)
selected_type = st.sidebar.selectbox("Type", ["All"] + types)
selected_lang = st.sidebar.selectbox("Language", ["All"] + languages)

# Show hint if nothing is uploaded
if not st.session_state.transcripts:
    st.sidebar.markdown("📁 Upload transcripts to enable meaningful filtering.")

# --- Session Management ---
st.sidebar.markdown("### 💾 Session Management")

# --- Save current session as JSON ---
json_data = json.dumps(st.session_state.transcripts, indent=2, ensure_ascii=False)
st.sidebar.download_button(
    "📤 Save Current Session (JSON)",
    json_data,
    file_name="transcripts_session.json",
    mime="application/json"
)

# --- Load a session from uploaded JSON file ---
if not st.session_state.get("clear_uploaded_json", False):
    uploaded_json = st.sidebar.file_uploader("📥 Load Session from JSON", type=["json"], key="load_json_key")

    if uploaded_json:
        try:
            loaded_transcripts = json.load(uploaded_json)
            if isinstance(loaded_transcripts, list):
                st.session_state["transcripts"] = loaded_transcripts
                for t in loaded_transcripts:
                    save_transcript_to_file(t)
                update_sidebar_filters()
                st.session_state["session_loaded"] = True
                st.session_state["clear_uploaded_json"] = True  # mark for clearing
                st.success("✅ Session loaded successfully!")
                st.rerun()
            else:
                st.sidebar.error("❌ Invalid JSON structure. Expected a list of transcripts.")
        except Exception as e:
            st.sidebar.error(f"❌ Failed to load session: {str(e)}")

# Reset the uploader after rerun
if st.session_state.get("clear_uploaded_json", False):
    st.session_state.pop("clear_uploaded_json")
    st.rerun()

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📄 Upload Transcript", 
    "📝 Overview", 
    "🔍 Search", 
    "🏷️ Tags & Themes", 
    "📤 Export"
])

with tab1:
    st.header("📄 Upload Transcript")

    if not st.session_state.get("clear_uploaded_transcript", False):
        uploaded_file = st.file_uploader("Upload a .txt or .docx file", type=["txt", "docx"], key="upload_transcript_key")

        if uploaded_file:
            if uploaded_file.type == "text/plain":
                text = StringIO(uploaded_file.getvalue().decode("utf-8")).read()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(uploaded_file)
            else:
                text = ""

            existing_locations = sorted({t["location"] for t in st.session_state.transcripts if t["location"]})
            existing_languages = sorted({t["language"] for t in st.session_state.transcripts if t["language"]})

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                location_option = st.selectbox("Research Site", existing_locations + ["Other"])
                location = st.text_input("Enter new Research Site") if location_option == "Other" else location_option
            with col2:
                t_type = st.selectbox("Type", ["FGD", "KII"])
            with col3:
                language_option = st.selectbox("Language", existing_languages + ["Other"])
                language = st.text_input("Enter new Language") if language_option == "Other" else language_option
            with col4:
                date_str = st.date_input("Date", datetime.now()).strftime("%Y-%m-%d")

            if st.button("Upload Transcript"):
                st.session_state.transcripts.insert(0, {
                    "text": text,
                    "location": location,
                    "type": t_type,
                    "date": date_str,
                    "language": language,
                    "tags": [],
                    "comments": [],
                    "themes": [],
                    "summary": ""
                })
                save_transcript_to_file(st.session_state.transcripts[0])
                update_sidebar_filters()
                st.session_state["clear_uploaded_transcript"] = True
                st.success("Transcript uploaded successfully!")
                st.rerun()

    # Show Generate DOCX button for all uploaded transcripts
    if st.session_state.transcripts:
        st.markdown("### 📥 Download DOCX for Uploaded Transcripts")
        for idx, transcript in enumerate(st.session_state.transcripts):
            docx_buffer = generate_docx(transcript)
            st.download_button(
                f"Generate DOCX for Transcript {idx+1} ({transcript['date']} • {transcript['type']} • {transcript['location']})",
                docx_buffer,
                file_name=f"{transcript['date']}_{transcript['type']}_{transcript['location'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"docx_download_{idx}"
            )

# Reset transcript uploader after rerun
if st.session_state.get("clear_uploaded_transcript", False):
    st.session_state.pop("clear_uploaded_transcript")
    st.rerun()

with tab2:
    st.header("📝 Transcript Overview")
    delete_index = None

    for idx, t in enumerate(st.session_state.transcripts):
        with st.expander(f"{t.get('date', '?')} • {t.get('type', '?')} • {t.get('location', '?')} • {t.get('language', 'Unknown')}"):
            st.markdown("### ✏️ Manual Summary and Themes")

            # Editable summary
            t["summary"] = st.text_area(f"Edit Summary (transcript {idx+1})", t.get("summary", ""), key=f"manual_summary_{idx}")

            # Editable themes
            themes_str = ", ".join(t.get("themes", []))
            new_themes_str = st.text_input(f"Edit Themes (comma-separated) (transcript {idx+1})", themes_str, key=f"manual_themes_{idx}")
            t["themes"] = [theme.strip() for theme in new_themes_str.split(",") if theme.strip()]

            save_transcript_to_file(t)

            st.markdown("### ✏️ Edit Transcript Info")
            orig_location = t.get("location", "")
            orig_type = t.get("type", "")
            orig_date = t.get("date", "")

            t["location"] = st.text_input(f"Edit Location (transcript {idx+1})", t["location"], key=f"edit_loc_{idx}")
            t["type"] = st.selectbox(
                f"Edit Type (transcript {idx+1})",
                ["FGD", "KII"],
                index=["FGD", "KII"].index(t.get("type", "FGD")),
                key=f"edit_type_{idx}"
            )
            t["language"] = st.text_input(f"Edit Language (transcript {idx+1})", t.get("language", ""), key=f"edit_lang_{idx}")
            t["date"] = st.date_input(
                f"Edit Date (transcript {idx+1})",
                datetime.strptime(t.get("date", "2025-01-01"), "%Y-%m-%d"),
                key=f"edit_date_{idx}"
            ).strftime("%Y-%m-%d")

            if (t["location"], t["type"], t["date"]) != (orig_location, orig_type, orig_date):
                delete_transcript_file({"location": orig_location, "type": orig_type, "date": orig_date})
            save_transcript_to_file(t)

            if st.button(f"🗑 Delete transcript {idx+1}"):
                delete_index = idx

    if delete_index is not None:
        transcript = st.session_state.transcripts.pop(delete_index)
        delete_transcript_file(transcript)
        st.rerun()

    # Visualization
    st.subheader("📊 Most Frequent Themes")
    theme_counter = Counter(theme for t in st.session_state.transcripts for theme in t.get("themes", []))
    if theme_counter:
        labels, values = zip(*theme_counter.most_common(10))
        fig, ax = plt.subplots()
        ax.barh(labels, values)
        ax.invert_yaxis()
        st.pyplot(fig)

with tab3:
    st.header("🔍 Search Transcripts")
    search_term = st.text_input("Search by keyword or phrase (e.g., discipline, barangay program)")
    def extract_context(text, keyword, context_lines=2):
        lines = text.split("\n")
        return ["\n".join(lines[max(i - context_lines, 0):min(i + context_lines + 1, len(lines))])
                for i, line in enumerate(lines) if keyword.lower() in line.lower()]
    # Use sidebar filters for search
    filtered = st.session_state.transcripts
    if selected_site != "All":
        filtered = [t for t in filtered if t["location"] == selected_site]
    if selected_type != "All":
        filtered = [t for t in filtered if t["type"] == selected_type]
    if selected_lang != "All":
        filtered = [t for t in filtered if t["language"] == selected_lang]

    # Additional filtering by tags and themes
    available_tags = sorted({tag for t in st.session_state.transcripts for tag in t["tags"]})
    available_themes = sorted({theme for t in st.session_state.transcripts for theme in t.get("themes", [])})

    selected_tags = st.multiselect("Filter by Tag", available_tags)
    selected_themes = st.multiselect("Filter by Theme", available_themes)

    # Apply tag and theme filters
    if selected_tags:
        filtered = [t for t in filtered if any(tag in t["tags"] for tag in selected_tags)]
    if selected_themes:
        filtered = [t for t in filtered if any(theme in t.get("themes", []) for theme in selected_themes)]

    st.subheader("📄 Matched Results")
    match_count = 0
    for idx, t in enumerate(filtered):
        if not search_term or search_term.lower() in t["text"].lower():
            snippets = extract_context(t["text"], search_term) if search_term else [t["text"]]
            match_count += len(snippets)
            with st.expander(f"{t['date']} • {t['type']} • {t['location']} • {t['language']} ({len(snippets)} match{'es' if len(snippets) > 1 else ''})"):
                for snip in snippets:
                    highlighted = re.sub(f"({re.escape(search_term)})", lambda m: f"**{m.group(1)}**", snip, flags=re.IGNORECASE)
                    st.markdown(f"🔍 {highlighted}")
                    st.markdown("---")
                comment = st.text_input(f"Add comment (transcript {idx+1})", key=f"comment_{idx}")
                tag = st.text_input(f"Add tag (transcript {idx+1})", key=f"tag_{idx}")
                if st.button(f"Save comment (transcript {idx+1})"):
                    t["comments"].append(comment)
                    save_transcript_to_file(t)
                    st.success("Comment added.")
                if st.button(f"Save tag (transcript {idx+1})"):
                    t["tags"].append(tag)
                    save_transcript_to_file(t)
                    st.success("Tag added.")
                if t["tags"]:
                    st.markdown("**Tags:** " + ", ".join(t["tags"]))
                if t["comments"]:
                    st.markdown("**Comments:**")
                    for c in t["comments"]:
                        st.markdown(f"- {c}")
    if search_term:
        st.info(f"🔍 Total matches found: {match_count}")

with tab4:
    st.header("🏷️ Tag and Theme Overview")
    tag_counts = Counter(tag for t in st.session_state.transcripts for tag in t["tags"])
    if tag_counts:
        st.subheader("Top Tags")
        tag_labels, tag_values = zip(*tag_counts.most_common(10))
        fig1, ax1 = plt.subplots()
        ax1.barh(tag_labels, tag_values)
        ax1.invert_yaxis()
        st.pyplot(fig1)

    # --- Theme Co-occurrence Heatmap ---
    st.subheader("📈 Theme Co-occurrence Heatmap")

    # Gather all unique themes
    theme_sets = [set(t.get("themes", [])) for t in st.session_state.transcripts if t.get("themes")]
    all_themes = sorted(set(theme for themes in theme_sets for theme in themes))

    if all_themes:
        # Initialize co-occurrence matrix as a DataFrame
        import pandas as pd
        co_df = pd.DataFrame(0, index=pd.Index(all_themes), columns=pd.Index(all_themes))

        for themes in theme_sets:
            for ti in themes:
                for tj in themes:
                    co_df.loc[ti, tj] += 1

        # Plot using seaborn
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        sns.heatmap(co_df, cmap="Blues", annot=False, fmt="d", ax=ax2)
        ax2.set_title("Theme Co-occurrence Matrix")
        st.pyplot(fig2)
    else:
        st.info("ℹ️ Not enough theme data to generate co-occurrence heatmap.")

with tab5:
    st.header("📤 Export Transcripts")
    col_export1, col_export2 = st.columns(2)

    with col_export1:
        json_data = json.dumps(st.session_state.transcripts, indent=2)
        st.download_button(
            "📤 Export All as JSON",
            json_data,
            file_name="transcripts.json",
            mime="application/json"
        )

    with col_export2:
        # Proper CSV escaping using .format() to avoid backslash in f-string
        csv_data = "text,location,type,language,date,tags,comments,themes,summary\n"
        for t in st.session_state.transcripts:
            text_escaped = t["text"].replace('"', '""')
            summary_escaped = t.get("summary", "").replace('"', '""')
            tags = ",".join(t["tags"])
            comments = ",".join(t["comments"])
            themes = ",".join(t.get("themes", []))

            line = '"{}",{},{},{},{},\"{}\",\"{}\",\"{}\",\"{}\"\n'.format(
                text_escaped,
                t["location"],
                t["type"],
                t["language"],
                t["date"],
                tags,
                comments,
                themes,
                summary_escaped
            )
            csv_data += line

        st.download_button(
            "📤 Export All as CSV",
            csv_data,
            file_name="transcripts.csv",
            mime="text/csv"
        )

# --- Merge two JSON files into current session ---
st.sidebar.markdown("### 🔗 Merge Two JSON Files")
merge_file1 = st.sidebar.file_uploader("Select First JSON File", type=["json"], key="merge_json1")
merge_file2 = st.sidebar.file_uploader("Select Second JSON File", type=["json"], key="merge_json2")

def merge_transcript_lists(list1, list2):
    """
    Merge two lists of transcripts, avoiding exact duplicates.
    Duplicates are detected by matching all fields.
    """
    existing = {json.dumps(t, sort_keys=True) for t in list1}
    merged = list1[:]
    for t in list2:
        t_json = json.dumps(t, sort_keys=True)
        if t_json not in existing:
            merged.append(t)
            existing.add(t_json)
    return merged

if merge_file1 and merge_file2:
    try:
        transcripts1 = json.load(merge_file1)
        transcripts2 = json.load(merge_file2)
        if isinstance(transcripts1, list) and isinstance(transcripts2, list):
            if st.sidebar.button("Merge Both Files Into Session"):
                merged = merge_transcript_lists(transcripts1, transcripts2)
                # Merge with current session transcripts
                st.session_state["transcripts"] = merge_transcript_lists(st.session_state["transcripts"], merged)
                # Save all to disk
                for t in st.session_state["transcripts"]:
                    save_transcript_to_file(t)
                update_sidebar_filters()
                st.sidebar.success("✅ Files merged into session!")
                st.rerun()
        else:
            st.sidebar.error("❌ Both files must be JSON lists of transcripts.")
    except Exception as e:
        st.sidebar.error(f"❌ Failed to merge: {str(e)}")